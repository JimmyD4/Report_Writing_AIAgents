import streamlit as st
import os
import shutil
import json
import traceback
import io 
from docx import Document 

# LangChain Imports
from langchain_groq import ChatGroq
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from dotenv import load_dotenv
from langchain_groq import ChatGroq


# --- Configuration ---
load_dotenv() # Load environment variables from .env file
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

CHROMA_DB_DIR = "./chroma_db_community_dev" # Directory to store your vector DB
PDF_SOURCE_DIR = "./pdfs_to_process" # Directory where you put your lecture PDFs
print("API keys loaded successfully!")
print(f"ChromaDB will be stored in: {CHROMA_DB_DIR}")
print(f"PDFs will be loaded from: {PDF_SOURCE_DIR}")


# --- Initialize Core Components ---
@st.cache_resource
def get_llm(llm_choice, groq_api_key, openai_api_key):
    """Initializes and caches the Language Model based on user choice."""
    if llm_choice == "Groq (Llama3-70b)":
        return ChatGroq(model_name="llama3-70b-8192", temperature=0.3, api_key=groq_api_key)
    elif llm_choice == "OpenAI (GPT-4o-mini)":
        return ChatOpenAI(model="gpt-4o-mini", temperature=0.3, api_key=openai_api_key)
    else:
        st.error("Invalid LLM choice. Please select either 'Groq (Llama3-70b)' or 'OpenAI (GPT-4o-mini)'.")
        st.stop()

@st.cache_resource
def get_embeddings(openai_api_key):
    """Initializes and caches the Embedding Model. OpenAIEmbeddings is used for consistency and quality."""
    return OpenAIEmbeddings(api_key=openai_api_key)

@st.cache_resource
def get_tavily_search_tool(tavily_api_key):
    """Initializes and caches the Tavily Search Tool."""
    return TavilySearchResults(max_results=5, tavily_api_key=tavily_api_key)

embeddings_model = get_embeddings(OPENAI_API_KEY)
tavily_search_tool = get_tavily_search_tool(TAVILY_API_KEY)


# --- PDF Processing and Vector Store Management ---
def process_and_store_pdfs(uploaded_files):
    """
    Processes uploaded PDF files, extracts text, chunks it, and stores embeddings
    in a ChromaDB vector store. Returns a ChromaDB retriever or None.
    """
    if os.path.exists(CHROMA_DB_DIR):
        st.info("Clearing existing lecture notes database...")
        shutil.rmtree(CHROMA_DB_DIR)
        st.info("Cleared.")

    all_docs = []
    
    if not uploaded_files:
        st.warning("No PDF files uploaded. The agent will rely solely on online research.")
        return None 

    temp_pdf_dir = "./temp_uploaded_pdfs"
    os.makedirs(temp_pdf_dir, exist_ok=True)

    for uploaded_file in uploaded_files:
        temp_file_path = os.path.join(temp_pdf_dir, uploaded_file.name)
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        st.write(f"Loading {uploaded_file.name}...")
        try:
            loader = PyPDFLoader(temp_file_path)
            docs = loader.load()
            all_docs.extend(docs)
        except Exception as e:
            st.error(f"ERROR: Could not load {uploaded_file.name}: {e}")
            continue
        finally:
            os.remove(temp_file_path)

    if os.path.exists(temp_pdf_dir):
        shutil.rmtree(temp_pdf_dir)

    if not all_docs:
        st.warning("No text extracted from PDFs. Please ensure they are not image-only PDFs or are valid PDF documents.")
        return None

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_documents(all_docs)
    st.write(f"Split {len(all_docs)} documents into {len(chunks)} chunks.")

    vectorstore = Chroma.from_documents(chunks, embeddings_model, persist_directory=CHROMA_DB_DIR)
    st.success(f"Successfully processed {len(uploaded_files)} PDF(s) and stored in vector DB.")
    return vectorstore.as_retriever()


# --- Agentic Logic (The AI's Workflow) ---
def conduct_research_and_report(user_prompt, lecture_retriever, llm_instance):
    """
    Main function to orchestrate the research, drafting, evaluation, and reporting process.
    """
    st.subheader("1. Gathering Information...")
    
    # 1. Researcher Agent (Online Research)
    st.info("Searching online for latest research papers...")
    research_query = f"Academic research papers on {user_prompt} in community development, published in last 10 years, include key ideas, findings and methodologies."
    online_research_results = tavily_search_tool.invoke({"query": research_query})
    
    with st.expander("Show Raw Online Research Findings (for debug)"):
        st.json(online_research_results)
    
    # Summarize online research results for context window efficiency
    st.info("Summarizing online research findings for context window efficiency...")
    research_summary_prompt = ChatPromptTemplate.from_messages([
        ("system", "You are a helpful assistant. Summarize the following online search results concisely, focusing on key findings, methodologies, and relevant conclusions for an academic report. Use bullet points or a brief paragraph per result. Prioritize information directly related to the user's main prompt. Keep the overall summary under 1000 words."),
        ("user", "Search Results:\n{results}")
    ])
    online_research_summary = (research_summary_prompt | llm_instance).invoke({"results": json.dumps(online_research_results)}).content
    st.write("### Summarized Online Research:")
    st.write(online_research_summary)

    # 2. Lecture Analyst Agent (Retrieve Lecture Context)
    st.info("Retrieving relevant information from your lecture slides...")
    
    lecture_context_str = ""
    if lecture_retriever:
        lecture_context_docs = lecture_retriever.get_relevant_documents(user_prompt)
        lecture_context_str = "\n".join([doc.page_content for doc in lecture_context_docs])
        if not lecture_context_str: 
            lecture_context_str = "No specific lecture content found relevant to the prompt within the processed lecture notes."
    else: 
        lecture_context_str = "No lecture context available as no PDFs were processed."

    st.write("### Relevant Lecture Notes:")
    st.write(lecture_context_str[:1500] + "..." if len(lecture_context_str) > 1500 else lecture_context_str)

    # --- Stage 3: Initial Drafting Agent ---
    st.subheader("2. Generating Initial Draft...")
    
    drafting_prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a top performing student and an expert academic writer in Community Development. 
         Your task is to write an initial draft of a report or literature review based on the user's prompt, 
         incorporating information from both online research findings and provided lecture notes.
         Ensure the draft is structured, academic in tone, and directly addresses the user's request.
         
         Summarized Online Research Findings:
         {online_research_summary}
         
         Lecture Notes Context:
         {lecture_notes}
         """),
        ("user", "Based on my prompt: '{user_prompt}', write an initial academic draft. "
        "Focus on synthesizing online research with concepts from the lecture notes."),
    ])
    
    initial_draft_chain = drafting_prompt | llm_instance
    initial_draft = initial_draft_chain.invoke({
        "online_research_summary": online_research_summary, 
        "lecture_notes": lecture_context_str,
        "user_prompt": user_prompt
    }).content
    
    st.write("### Initial Draft:")
    st.write(initial_draft)

    # --- Stage 4: Critic/Evaluator Agent ---
    st.subheader("3. Evaluating the Draft...")
    critic_prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a critical academic peer reviewer specializing in Community Development. 
         Your task is to evaluate the provided 'initial draft' based on the 'original user prompt' and the 'lecture context'.
         Provide constructive feedback using the following criteria:
         -   **Accuracy & Completeness:** Is the information factual and does it fully address the original prompt? Are there any obvious gaps?
         -   **Relevance to Lecture Notes:** Does it effectively integrate and reference concepts, theories, or examples from the provided lecture notes? Are there missed opportunities for connection?
         -   **Integration of Online Research:** Does it properly synthesize and conceptually reference the online research findings (from the summarized content provided)? Is it up-to-date?
         -   **Academic Tone & Structure:** Is the language appropriate for a postgraduate academic paper? Is the structure logical and easy to follow?
         -   **Referencing:** Check if the report has reference and also if the online papers has been quoted reference in the report in APA style format. If not, suggest how to add them.
         -   **AI detection:** Does the report look like it's written by a human. What is the percentage of the writeup that looks like it's written by an AI and how can we humanize it? 
         
         Provide your feedback in bullet points, clearly indicating areas for improvement. Do NOT rewrite the draft yourself.
         """),
        ("user", "Original User Prompt: '{original_user_prompt}'\n\nLecture Notes Context: '{lecture_context}'\n\nInitial Draft to Evaluate:\n'{initial_draft}'\n\nSummarized Online Research for reference: {online_research_summary}") 
    ])
    
    critic_chain = critic_prompt | llm_instance
    critic_feedback = critic_chain.invoke({
        "original_user_prompt": user_prompt,
        "lecture_context": lecture_context_str,
        "initial_draft": initial_draft,
        "online_research_summary": online_research_summary 
    }).content
    
    st.write("### Critic's Feedback:")
    st.write(critic_feedback)

    # --- Stage 5: Final Report Agent ---
    st.subheader("4. Generating Final Report...")
    final_report_prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a highly skilled academic writer for a Masters in Community Development. 
         Your task is to revise and finalize the 'initial draft' based on the 'critic's feedback', the 'original user prompt', 
         the 'online research findings', and the 'lecture notes context'.
         Produce a polished, comprehensive, and academically sound report that integrates all sources seamlessly.
         
         Original User Prompt: '{original_user_prompt}'
         
         Summarized Online Research Findings:
         {online_research_summary}
         
         Lecture Notes Context:
         {lecture_notes}
         
         Initial Draft:
         {initial_draft}
         
         Critic's Feedback:
         {critic_feedback}
         
         Produce the final, refined report.
         """),
        ("user", "Please finalize the report, incorporating the feedback and all provided context.")
    ])
    
    final_report_chain = final_report_prompt | llm_instance
    final_report = final_report_chain.invoke({
        "original_user_prompt": user_prompt,
        "online_research_summary": online_research_summary, 
        "lecture_notes": lecture_context_str,
        "initial_draft": initial_draft,
        "critic_feedback": critic_feedback
    }).content
    
    st.subheader("🎉 Final Report:")
    st.write(final_report)

    # --- NEW FEATURE: Download Word Document ---
    st.markdown("---")
    st.subheader("Download Report")

    # Create a new Word Document
    document = Document()
    document.add_heading("Community Development Research Report", level=1) # Add a title
    
    # Add the final report content as a single paragraph (you could add more formatting here)
    document.add_paragraph(final_report) 

    # Save document to a BytesIO object (in memory)
    byte_io = io.BytesIO()
    document.save(byte_io)
    byte_io.seek(0) # Rewind the buffer to the beginning

    # Store the byte_io object in session state to make it available for the download button
    st.session_state['download_report_data'] = byte_io.getvalue()
    st.session_state['report_generated'] = True # Flag to show the button

# --- Streamlit UI ---
st.set_page_config(page_title="Community Dev Research Agent", layout="wide")
st.title("👨‍🎓 Community Development Research Agent")
st.markdown("Your AI assistant for Masters research, integrating lecture notes and latest online papers.")

# Sidebar for LLM choice and PDF upload
st.sidebar.header("Configuration")

llm_choice = st.sidebar.selectbox(
    "Choose your Large Language Model:",
    ("Groq (Llama3-70b)", "OpenAI (GPT-4o-mini)")
)

current_llm = get_llm(llm_choice, GROQ_API_KEY, OPENAI_API_KEY)


st.sidebar.header("Upload Lecture Slides (PDFs)")
uploaded_files = st.sidebar.file_uploader(
    "Upload your lecture PDFs here (multiple files accepted)",
    type="pdf",
    accept_multiple_files=True
)

if "lecture_retriever" not in st.session_state:
    st.session_state.lecture_retriever = None
# Reset download state when new PDFs are uploaded or if files change
if "report_generated" not in st.session_state:
    st.session_state.report_generated = False


if uploaded_files:
    if st.session_state.get('last_uploaded_file_names') != [f.name for f in uploaded_files] or st.sidebar.button("Re-process PDFs"):
        st.session_state.last_uploaded_file_names = [f.name for f in uploaded_files]
        with st.spinner("Processing PDFs and building knowledge base..."):
            st.session_state.lecture_retriever = process_and_store_pdfs(uploaded_files)
        if st.session_state.lecture_retriever:
            st.sidebar.success("PDFs processed and ready!")
        else:
            st.sidebar.info("No lecture knowledge base created (no text extracted or no valid PDFs).")
        st.session_state.report_generated = False # Reset report generated flag on new PDF processing
    else:
        st.sidebar.info("PDFs loaded from previous session or already processed.")
        
    st.info(f"{len(uploaded_files)} PDF(s) loaded.")
else:
    st.session_state.lecture_retriever = None
    st.sidebar.info("No PDFs currently uploaded. The agent will rely solely on online research.")
    # If no files are uploaded, clear the last_uploaded_file_names to ensure re-processing on next upload
    if 'last_uploaded_file_names' in st.session_state:
        del st.session_state['last_uploaded_file_names']


user_prompt = st.text_area(
    "Enter your research prompt or question:",
    value="Write a comprehensive report on the role of participatory approaches in sustainable community development, "
          "incorporating findings from recent empirical studies (last 5-10 years).",
    height=150
)

if st.button("Generate Research Report"):
    st.session_state.report_generated = False # Reset flag before new generation attempt
    st.session_state.download_report_data = None # Clear previous download data
    if not user_prompt:
        st.error("Please enter a research prompt.")
    else:
        with st.spinner("🚀 Generating report... This might take a few moments."):
            try:
                conduct_research_and_report(user_prompt, st.session_state.lecture_retriever, current_llm)
            except Exception as e:
                st.error(f"An error occurred: {e}. Please check your API keys and prompt.")
                st.exception(e) # This will print the full traceback in Streamlit's logs for debugging

# --- NEW FEATURE: Display Download Button ---
if st.session_state.report_generated and st.session_state.download_report_data:
    st.download_button(
        label="Download Final Report as Word Document",
        data=st.session_state.download_report_data,
        file_name="community_development_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        help="Click to download the generated report as a Microsoft Word document."
    )
# --- END NEW FEATURE ---

st.markdown("---")
st.caption("Disclaimer: This AI agent is a tool to assist your research. Always critically review generated content and verify information for accuracy and academic integrity. Citing sources properly is crucial.")